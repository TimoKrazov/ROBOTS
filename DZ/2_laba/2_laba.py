import pandas as pd
from openpyxl import load_workbook
from openpyxl.chart import BarChart, Reference


#----Excel----
csv_file = pd.read_csv('Supermarket_Sales.csv')
csv_file['Стоимость'] = csv_file['Unit price'] * csv_file['Quantity']
csv_file.to_excel("Raport.xlsx", sheet_name="Закупки", index= False)


wb = load_workbook('Raport.xlsx')
ws = wb['Закупки']

barChart = BarChart()
barChart.type = "col"
barChart.title = "Стоимость товаров"
barChart.y_axis.title = "Стоимость"

data_pr = Reference(ws, min_col=csv_file.columns.get_loc('Стоимость') + 1,min_row= 2, max_row=len(csv_file)+1)
name_pr = Reference(ws, min_col=csv_file.columns.get_loc('Invoice ID') + 1,min_row= 2, max_row=len(csv_file)+1)

barChart.add_data(data_pr)
barChart.set_categories(name_pr)
ws.add_chart(barChart, "S1")
wb.save('Raport.xlsx')



#---Word---
from docx import Document


doc = Document()
doc.add_heading('Автоматизированный отчёт')
doc.add_paragraph('Отчет сгенерирован роботом...')

table = doc.add_table(rows=1, cols = 4)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'ID'
table.rows[0].cells[1].text = 'Количество'
table.rows[0].cells[2].text = 'Цена'
table.rows[0].cells[3].text = 'Стоимость'

#---для будующего pdf---
num_rows = len(csv_file)
summary = 0

#-------------------
for _, row in csv_file.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = str(row['Invoice ID'])
    row_cells[1].text = str(row['Unit price'])
    row_cells[2].text = str(row['Quantity'])
    row_cells[3].text = str(row['Стоимость'])
    summary += row['Стоимость']
doc.save('Отчётик.docx')



#---PDF---
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch #Дюймы
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

pdfmetrics.registerFont(TTFont('Arial', 'C:/Windows/Fonts/Arial.ttf'))


doc = SimpleDocTemplate("Отчётик.pdf", pagesize=A4)

elements = []
styles = getSampleStyleSheet()

logo = Image("logo.png", width=2*inch, height=1*inch)

elements.append(logo)
elements.append(Spacer(1, 12))

text_style = ParagraphStyle(name = 'text', fontName='Arial', fontSize=14)

title = Paragraph("Итоговая сводка", text_style)

elements.append(title)
elements.append(Spacer(1,12))

text = f'''
Общая стоимость всех позиций: {summary:,.2f} руб. <br/>
Количество позиций: {num_rows} <br/> 
'''

text_paragraph = Paragraph(text, text_style)

elements.append(text_paragraph)


doc.build(elements)



#---И ZIP---
from zipfile import ZipFile
from datetime import datetime

with ZipFile(f'Отчёт_{datetime.now().strftime("%Y-%m-%d")}.zip', 'w') as zip:
    for file in ["Raport.xlsx", "Отчётик.pdf", "Отчётик.docx"]:
        zip.write(file)

print('ВСЁ')